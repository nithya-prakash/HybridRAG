from io import BytesIO

import docx

from app.services.parsing.errors import ParsingError
from app.services.parsing.models import BLOCK_TYPE_PARAGRAPH, ParsedBlock, ParsedDocument
from app.services.parsing.page_estimate import estimate_page_number

_HEADING_STYLE_PREFIX = "Heading "
_TITLE_STYLE_NAME = "Title"


def _heading_level(style_name: str) -> int | None:
    if style_name == _TITLE_STYLE_NAME:
        return 1
    if style_name.startswith(_HEADING_STYLE_PREFIX):
        suffix = style_name[len(_HEADING_STYLE_PREFIX) :]
        if suffix.isdigit():
            return int(suffix)
    return None


def parse_docx(content: bytes) -> ParsedDocument:
    try:
        document = docx.Document(BytesIO(content))
        paragraphs = list(document.paragraphs)
    except Exception as exc:
        # python-docx has no single exception base class for "not a valid docx" —
        # a corrupt zip, a zip missing the expected OPC parts, and a malformed
        # document.xml all surface as different exception types.
        raise ParsingError(f"Could not parse DOCX: {exc}") from exc

    blocks: list[ParsedBlock] = []
    heading_stack: list[tuple[int, str]] = []
    offset = 0

    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        level = _heading_level(paragraph.style.name if paragraph.style else "")
        if level is not None:
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text))
            offset += len(text) + 2
            continue

        blocks.append(
            ParsedBlock(
                text=text,
                block_type=BLOCK_TYPE_PARAGRAPH,
                heading_path=[title for _, title in heading_stack],
                page_number=estimate_page_number(offset),
            )
        )
        offset += len(text) + 2

    return ParsedDocument(blocks=blocks)
